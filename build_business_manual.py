from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path(r"D:\WorkSpace\Code\novel\交付物\小说自动生产与发布系统_业务操作与维护指导书_V1.0.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

BLUE = "2E74B5"
DARK = "1F4D78"
NAVY = "17365D"
LIGHT = "E8EEF5"
PALE = "F4F6F9"
GREEN = "E2F0D9"
YELLOW = "FFF2CC"
RED = "FCE4D6"
GRAY = "666666"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            set_cell_margins(row.cells[idx])
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, size=None, bold=None, color=None, italic=None, name="Microsoft YaHei"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.italic = italic


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_font(run, 9, color=GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    run = paragraph.add_run(" 页")
    set_font(run, 9, color=GRAY)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_font(r, 11, bold=True, color=NAVY)
        r = p.add_run(text[len(bold_prefix):])
        set_font(r, 11)
    else:
        r = p.add_run(text)
        set_font(r, 11)
    return p


def add_bullets(doc, items, checked=False):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        if checked:
            p.add_run("□ ")
        r = p.add_run(item)
        set_font(r, 11)


def add_steps(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_font(r, 11)


def add_callout(doc, title, body, fill=PALE, color=NAVY):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    prevent_row_split(table.rows[0])
    set_table_widths(table, [6.25])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_font(r, 11, bold=True, color=color)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(body)
    set_font(r, 10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths, status_col=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0])
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        set_cell_shading(c, LIGHT)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, 10.5, bold=True, color=NAVY)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            if i == 0 or (status_col is not None and i == status_col):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            set_font(r, 10)
            if status_col is not None and i == status_col:
                value_s = str(value)
                if "正常" in value_s or "成功" in value_s or "通过" in value_s:
                    set_cell_shading(cells[i], GREEN)
                elif "等待" in value_s or "审核" in value_s or "跳过" in value_s:
                    set_cell_shading(cells[i], YELLOW)
                elif "异常" in value_s or "失败" in value_s or "离线" in value_s:
                    set_cell_shading(cells[i], RED)
    set_table_widths(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.9)
section.right_margin = Inches(0.9)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Microsoft YaHei"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, DARK, 10, 5),
):
    st = styles[name]
    st.font.name = "Microsoft YaHei"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

for name in ("List Bullet", "List Number"):
    st = styles[name]
    st.font.name = "Microsoft YaHei"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    st.font.size = Pt(11)
    st.paragraph_format.left_indent = Inches(0.38)
    st.paragraph_format.first_line_indent = Inches(-0.19)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.line_spacing = 1.25

header = section.header
p = header.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("小说自动生产与发布系统｜业务操作与维护指导书")
set_font(r, 8.5, color=GRAY)
add_page_number(section.footer.paragraphs[0])

# Cover
doc.add_paragraph().paragraph_format.space_after = Pt(48)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("业务操作与维护指导书")
set_font(r, 28, bold=True, color=NAVY)
p.paragraph_format.space_after = Pt(8)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("小说自动生产与发布系统")
set_font(r, 18, color=BLUE)
p.paragraph_format.space_after = Pt(36)

add_callout(doc, "适用对象", "小说运营人员、内容审核人员、值班人员和业务负责人。本文档不要求使用者具备编程经验。", fill=LIGHT)
doc.add_paragraph().paragraph_format.space_after = Pt(16)
add_table(doc, ["文档信息", "内容"], [
    ("版本", "V1.0"),
    ("发布日期", "2026年8月"),
    ("运行环境", "已完成迁移并通过测试的新电脑"),
    ("适用范围", "五部小说的章节生产、审核、发布、状态跟踪与日常维护"),
], [1.6, 4.65])
doc.add_paragraph().paragraph_format.space_after = Pt(18)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("使用原则：先检查、再运行；先核实、再重试；异常留痕、不重复发布。")
set_font(r, 11, bold=True, color=DARK)
doc.add_page_break()

# 1
add_heading(doc, "1. 文档目的与系统概览", 1)
add_body(doc, "本系统用于将飞书中的小说和章节任务转化为可执行工作，并通过五台云手机上的番茄作家助手完成章节发布。业务人员主要负责准备任务、完成审核、观察运行状态、处理常见异常以及做好交接记录。")
add_heading(doc, "1.1 系统完成的工作", 2)
add_steps(doc, [
    "读取飞书中的小说、章节、账号、排期和审核状态。",
    "按照规则生成或筛选可发布的章节。",
    "通过 ADB 控制对应云手机，进入番茄作家助手。",
    "填写章节标题和正文，发布时选择“有使用AI”。",
    "检查发布结果并把状态回写到飞书。",
    "发布成功或出现异常后重置页面，为下一次运行做准备。",
])
add_heading(doc, "1.2 业务人员需要关注什么", 2)
add_bullets(doc, [
    "飞书中的任务是否完整、审核状态是否正确。",
    "五台云手机是否在线、账号是否登录、应用是否正常。",
    "运行日志是否持续产生，以及是否出现重复异常。",
    "飞书状态、手机端章节列表与实际发布结果是否一致。",
])
add_callout(doc, "重要说明", "程序会跳过未达到发布条件的章节。章节被跳过通常代表门禁条件未满足，不等同于系统故障。", fill=YELLOW)

# 2
add_heading(doc, "2. 角色与职责", 1)
add_table(doc, ["角色", "主要职责", "不应执行的操作"], [
    ("内容运营", "维护小说、章节、排期和库存；发起生产或发布任务。", "直接修改程序文件或删除运行数据。"),
    ("内容审核", "检查标题、正文、敏感内容和完整性；确认是否可定稿。", "未查看正文就批量标记审核通过。"),
    ("值班人员", "检查设备、服务、日志和异常；做好交接记录。", "未核实平台结果就重复点击发布。"),
    ("业务负责人", "确认发布策略、异常升级和暂停/恢复决定。", "绕过门禁批量发布未审核内容。"),
    ("技术维护", "处理配置、依赖、设备连接、坐标和程序故障。", "在业务运行中无通知地重启或更换设备绑定。"),
], [1.0, 3.1, 2.15])

# 3
add_heading(doc, "3. 每日标准操作流程", 1)
add_heading(doc, "3.1 班前检查（建议每天首次运行前）", 2)
add_bullets(doc, [
    "确认新电脑已经开机、网络正常、系统时间正确。",
    "确认红手指云手机客户端正常，五台云手机均已开启。",
    "确认五台手机中的番茄作家助手账号保持登录。",
    "确认飞书中待处理章节的标题、正文、审核和排期信息完整。",
    "确认上一班次没有未处理的发布异常或重复发布风险。",
], checked=True)
add_heading(doc, "3.2 准备章节", 2)
add_steps(doc, [
    "在飞书中选择对应小说和章节。",
    "确认章节名称、正文、最终字数和当前最终版本正确。",
    "完成内容审核，并把生产状态设置为已定稿、已审核或已完成。",
    "将内容锁定状态设置为“是”或“人工锁定”。",
    "确认发布状态为“待发布”，并检查计划发布时间。",
    "确认小说自动发布开关已开启，账号健康状态正常。",
])
add_heading(doc, "3.3 启动或确认无人值守运行", 2)
add_body(doc, "正常情况下，电脑登录后系统会自动启动。业务人员不需要反复打开程序。确认运行状态时，可查看服务日志是否持续更新，或请技术人员检查后台服务。")
add_callout(doc, "不要重复启动", "发现程序已经运行时，不要再次双击启动脚本。多个实例可能争用同一台手机，并造成状态混乱。", fill=RED, color="9B1C1C")
add_heading(doc, "3.4 运行后核对", 2)
add_bullets(doc, [
    "发布成功的章节在番茄作家助手中可见。",
    "章节标题和正文完整，没有只填写标题、正文为空的情况。",
    "飞书发布状态已经更新，错误信息为空或与实际情况一致。",
    "手机页面已恢复到后续任务可继续执行的状态。",
])

# 4
add_heading(doc, "4. 五台设备与小说对应关系", 1)
add_table(doc, ["账号编号", "小说名称", "业务核对重点"], [
    ("NOVEL-001", "灵气复苏", "登录账号、书名与章节列表一致"),
    ("NOVEL-002", "重生归来", "登录账号、书名与章节列表一致"),
    ("NOVEL-003", "世界突变", "登录账号、书名与章节列表一致"),
    ("NOVEL-004", "守夜人觉醒", "登录账号、书名与章节列表一致"),
    ("NOVEL-005", "传承戒指", "登录账号、书名与章节列表一致"),
], [1.35, 1.55, 3.35])
add_callout(doc, "绑定纪律", "设备、账号和小说必须一一对应。更换端口不代表可以更换账号顺序；任何绑定调整都应由技术维护人员完成并记录。", fill=YELLOW)
add_heading(doc, "4.1 手机端应保持的状态", 2)
add_bullets(doc, [
    "无锁屏、系统升级、权限申请和其他遮挡弹窗。",
    "番茄作家助手可以正常联网，账号未退出。",
    "不要同时在其他手机上操作同一个作者账号。",
    "分辨率和显示缩放不要随意改变。",
])

# 5
add_heading(doc, "5. 业务状态识别指南", 1)
add_table(doc, ["状态", "业务含义", "应该怎么做"], [
    ("待生产", "等待系统生成正文。", "保持任务信息完整，等待下一周期。"),
    ("生产中", "系统正在处理章节。", "避免修改正文或重复发起。"),
    ("待审核", "正文已生成，等待人工检查。", "审核标题、内容、字数和连续性。"),
    ("已定稿/已完成", "内容已确认，可进入发布判断。", "检查锁定、排期和发布状态。"),
    ("待发布", "满足基础状态，等待门禁和排期。", "确认账号、库存和计划时间。"),
    ("审核中", "已经提交平台，等待结果刷新。", "进入审核点刷新状态，避免重复提交。"),
    ("已发布", "平台已确认发布成功。", "核对章节列表并归档。"),
    ("发布异常", "设备、页面或网络等环节失败。", "查看错误信息，按第7章处理。"),
], [1.25, 2.2, 2.8], status_col=0)
add_heading(doc, "5.1 为什么任务会被跳过", 2)
add_bullets(doc, [
    "人工审核尚未完成。",
    "正文没有锁定或缺少最终版本。",
    "还没有到计划发布时间。",
    "库存低于安全门槛。",
    "小说自动发布开关关闭。",
    "账号处于养号期或健康状态异常。",
    "章节已经发布，系统通过幂等机制避免重复处理。",
])

# 6
add_heading(doc, "6. 发布成功判定与业务复核", 1)
add_body(doc, "发布按钮被点击并不等于发布成功。业务人员应以平台章节列表和飞书状态一致为准。")
add_heading(doc, "6.1 成功判定标准", 2)
add_bullets(doc, [
    "番茄作家助手章节管理页面存在对应章节。",
    "章节标题正确，不出现问号、乱码或错书。",
    "正文完整，字数与飞书最终版本基本一致。",
    "平台状态显示已发布或审核通过。",
    "飞书发布状态更新为已发布，且没有遗留错误信息。",
], checked=True)
add_heading(doc, "6.2 审核状态刷新", 2)
add_body(doc, "平台审核状态有时只有进入章节管理或审核页面后才会刷新。遇到飞书显示审核中、但业务人员已知平台已通过时，应先进入对应审核点刷新页面，再进行状态同步。")
add_callout(doc, "严禁重复发布", "状态不确定时先核实章节列表，不要直接把章节改回“待发布”。重复提交可能生成重复章节。", fill=RED, color="9B1C1C")

# 7
add_heading(doc, "7. 常见异常与处理办法", 1)
add_table(doc, ["现象", "可能原因", "业务人员处理"], [
    ("设备显示离线", "云手机关闭、端口变化或连接中断。", "确认云手机已开启，记录设备和时间，通知技术重新连接。"),
    ("只写了标题，正文为空", "正文输入失败、页面未进入编辑区或超时。", "立即停止该章节后续发布，保留现场和日志，通知技术处理。"),
    ("页面卡住或出现弹窗", "更新、权限、登录或系统弹窗遮挡。", "关闭弹窗，确认账号仍登录；必要时按HOME后重新打开应用。"),
    ("显示审核中很久", "平台状态尚未刷新。", "进入章节管理/审核页面刷新，再核对飞书。"),
    ("章节标题乱码", "字符输入或编码异常。", "不要继续批量发布；记录章节ID并升级处理。"),
    ("疑似重复发布", "状态回写失败、多实例运行或人工改回待发布。", "暂停对应小说，核对平台章节列表和飞书记录。"),
    ("连续多章失败", "设备、网络或应用版本发生整体变化。", "暂停无人值守发布，保留日志并通知技术负责人。"),
], [1.55, 2.15, 2.55])
add_heading(doc, "7.1 异常处理四步法", 2)
add_steps(doc, [
    "暂停：存在重复发布、正文为空或错书风险时，先暂停对应小说或整套发布。",
    "核实：查看手机章节列表、飞书状态和错误信息，判断平台是否已实际收录。",
    "留痕：记录时间、小说、章节ID、设备、现象和已采取的动作；必要时截图。",
    "升级：将记录交给技术维护人员，修复后先进行单章验证，再恢复批量运行。",
])

# 8
add_heading(doc, "8. 日常维护与巡检", 1)
add_heading(doc, "8.1 每日巡检", 2)
add_bullets(doc, [
    "检查五台云手机是否在线。",
    "检查作者账号是否保持登录。",
    "检查待发布、审核中和发布异常的章节数量。",
    "检查是否存在同一错误连续出现三次以上。",
    "抽查至少一章，确认飞书与平台状态一致。",
], checked=True)
add_heading(doc, "8.2 每周巡检", 2)
add_bullets(doc, [
    "统计本周生产数量、发布数量、成功率和主要异常。",
    "检查章节库存是否达到安全门槛。",
    "检查五台手机是否出现应用升级或页面变化。",
    "确认日志和本地数据有备份，磁盘空间充足。",
    "复核设备、账号、小说绑定表。",
], checked=True)
add_heading(doc, "8.3 维护边界", 2)
add_table(doc, ["业务人员可自行处理", "交给技术人员处理"], [
    ("检查飞书状态、刷新审核页、关闭普通弹窗、确认账号登录、记录异常。", "修改程序、修改坐标、修改ADB端口、重建环境、修复数据库、处理持续HTTP 500/503。"),
], [3.1, 3.15])

# 9
add_heading(doc, "9. 启停、暂停与恢复", 1)
add_heading(doc, "9.1 哪些情况应暂停发布", 2)
add_bullets(doc, [
    "出现错书、错账号或重复章节。",
    "章节正文为空、乱码或明显不完整。",
    "五台设备大面积离线。",
    "番茄作家助手升级后页面结构发生变化。",
    "飞书状态与平台结果持续不一致。",
    "业务负责人要求暂停。",
])
add_heading(doc, "9.2 恢复运行前确认", 2)
add_bullets(doc, [
    "异常原因已经明确并处理。",
    "设备在线，账号和小说对应正确。",
    "待发布章节状态没有被重复修改。",
    "先完成一章单次测试且发布结果正确。",
    "通知相关业务人员恢复时间。",
], checked=True)
add_callout(doc, "恢复原则", "先单章、后单周期、再无人值守。不要修复后立即批量恢复。", fill=LIGHT)

# 10
add_heading(doc, "10. 值班交接与异常记录", 1)
add_heading(doc, "10.1 交接内容", 2)
add_bullets(doc, [
    "当前服务是否运行。",
    "五台设备在线情况。",
    "待审核、待发布、审核中和发布异常数量。",
    "当班已经处理的异常及处理结果。",
    "仍待跟进的章节ID、设备和责任人。",
    "下一次需要检查的时间点。",
])
add_heading(doc, "10.2 异常记录模板", 2)
add_table(doc, ["记录项", "填写内容"], [
    ("发生时间", "年/月/日 时:分"),
    ("小说/章节", "小说名称、章节ID、章节名称"),
    ("设备", "账号编号或设备编号"),
    ("异常现象", "看到的页面、状态或错误信息"),
    ("平台实际结果", "未提交/审核中/已发布/不确定"),
    ("已采取动作", "刷新、重置页面、暂停、通知技术等"),
    ("处理结论", "已恢复/继续观察/等待技术/暂停发布"),
    ("责任人与复查时间", "姓名或岗位、计划复查时间"),
], [1.55, 4.7])

# 11
add_heading(doc, "11. 业务验收与周期复盘", 1)
add_heading(doc, "11.1 每批任务验收", 2)
add_bullets(doc, [
    "计划发布数量与实际发布数量一致。",
    "没有标题乱码、正文为空和错书问题。",
    "所有成功章节都已在飞书标记为已发布。",
    "失败章节均有错误信息和后续责任人。",
    "设备页面已复位，可以进入下一批任务。",
], checked=True)
add_heading(doc, "11.2 建议关注的业务指标", 2)
add_table(doc, ["指标", "计算口径", "用途"], [
    ("发布成功率", "成功发布章节数 ÷ 实际尝试章节数", "判断系统稳定性"),
    ("重复发布数", "同一作品出现重复章节的次数", "必须重点控制"),
    ("正文完整率", "标题与正文均完整的章节占比", "判断输入流程质量"),
    ("状态一致率", "飞书与平台状态一致的章节占比", "判断闭环质量"),
    ("人工介入次数", "需要人工处理的异常次数", "评估无人值守效果"),
], [1.45, 3.0, 1.8])

# 12 quick card
add_heading(doc, "12. 一页式值班速查", 1)
add_table(doc, ["时点", "必须完成的动作"], [
    ("上班后", "开电脑 → 看云手机 → 看飞书异常 → 确认服务运行"),
    ("发布前", "核对章节 → 完成人工审核 → 锁定正文 → 检查排期和账号"),
    ("发布后", "看平台章节 → 看标题正文 → 看飞书回写 → 看页面复位"),
    ("遇异常", "先暂停 → 再核实 → 做记录 → 通知责任人"),
    ("恢复时", "先单章 → 再单周期 → 最后恢复无人值守"),
    ("交班前", "写设备状态、任务状态、遗留异常和复查时间"),
], [1.25, 5.0])
add_callout(doc, "三个绝对不要", "不要把点击发布当作发布成功；不要在状态不确定时重复提交；不要在程序运行时随意修改设备、账号和小说绑定。", fill=RED, color="9B1C1C")

# Appendix
add_heading(doc, "附录A：需要技术协助时提供的信息", 1)
add_bullets(doc, [
    "异常发生的准确时间。",
    "小说名称、章节ID和章节名称。",
    "对应账号编号或云手机。",
    "飞书当前状态和错误信息。",
    "手机端实际页面或截图。",
    "是否已经刷新、返回HOME或重新打开应用。",
    "是否可能已经在平台发布。",
])
add_heading(doc, "附录B：技术维护入口（仅供指定人员）", 1)
add_body(doc, "项目目录：D:\\WorkSpace\\Code\\novel\\openclaw")
add_table(doc, ["用途", "命令或文件"], [
    ("健康检查", ".\\.venv\\Scripts\\python.exe .\\scripts\\healthcheck.py"),
    ("单周期验证", ".\\.venv\\Scripts\\python.exe .\\closed_loop.py --once"),
    ("持续运行", ".\\.venv\\Scripts\\python.exe .\\closed_loop.py --continuous"),
    ("后台启动脚本", ".\\scripts\\run_closed_loop.ps1"),
    ("业务日志", ".\\logs\\openclaw.log"),
    ("服务日志", ".\\logs\\closed_loop.service.log"),
    ("设备配置", ".\\config\\config.yaml"),
], [1.6, 4.65])
# prevent widows and apply font to all generated runs
for p in doc.paragraphs:
    p.paragraph_format.widow_control = True
    for run in p.runs:
        if not run.font.name:
            set_font(run, 11)

doc.core_properties.title = "小说自动生产与发布系统 业务操作与维护指导书"
doc.core_properties.subject = "面向业务人员的操作、巡检、异常处理与交接手册"
doc.core_properties.author = "项目交付组"
doc.save(OUT)
print(OUT)
